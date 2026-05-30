import os
import re
import shutil
import time
import platform
from pathlib import Path
from datetime import datetime
from collections import OrderedDict

try:
    from wcmatch import glob as wc_glob
    _WCMATCH = True
except ImportError:
    _WCMATCH = False

try:
    import humanize
    _HUMANIZE = True
except ImportError:
    _HUMANIZE = False

try:
    import git
    _GITPYTHON = True
except ImportError:
    _GITPYTHON = False

try:
    import magic
    _FILE_MAGIC = True
except ImportError:
    _FILE_MAGIC = False

try:
    from watchdog.observers import Observer
    from watchdog.events import FileSystemEventHandler
    _WATCHDOG = True
except ImportError:
    _WATCHDOG = False

try:
    import send2trash
    _SEND2TRASH = True
except ImportError:
    _SEND2TRASH = False

_OS = platform.system()  # "Windows" | "Darwin" | "Linux"

# Recent files cache — tracks files created/written by the assistant
_RECENT_FILES: OrderedDict[str, str] = OrderedDict()  # filename.lower() -> full path

# Lazy import for folder name index
_FOLDER_INDEX = None

def _get_folder_index():
    global _FOLDER_INDEX
    if _FOLDER_INDEX is None:
        from actions.folder_index import folder_index
        _FOLDER_INDEX = folder_index
    return _FOLDER_INDEX

def _track_file(path: Path):
    """Store file in recent files cache (LRU)."""
    resolved = str(path.resolve())
    _RECENT_FILES[path.name.lower()] = resolved
    _RECENT_FILES.move_to_end(path.name.lower())
    _RECENT_FILES[path.stem.lower()] = resolved
    _RECENT_FILES.move_to_end(path.stem.lower())
    # Evict oldest (LRU) when over limit
    while len(_RECENT_FILES) > 50:
        _RECENT_FILES.popitem(last=False)

def _find_recent(name: str) -> str | None:
    """Look up a file by name in recent files cache."""
    key = name.strip().lower()
    if key in _RECENT_FILES:
        return _RECENT_FILES[key]
    # Try partial match
    for k, v in _RECENT_FILES.items():
        if key in k or k in key:
            return v
    return None

_SAFE_ROOTS: list[Path] = [
    Path.home(),
]

if _OS == "Windows":
    import string
    for letter in string.ascii_uppercase:
        drive = f"{letter}:\\"
        if os.path.exists(drive):
            _SAFE_ROOTS.append(Path(drive))
elif _OS == "Linux" or _OS == "Darwin":
    _SAFE_ROOTS.append(Path("/"))

_BLOCKED_PATHS: list[Path] = []
if _OS == "Windows":
    windir = os.environ.get("SystemRoot", "C:\\Windows")
    _BLOCKED_PATHS = [
        Path(windir),
        Path("C:\\Program Files"),
        Path("C:\\Program Files (x86)"),
        Path("C:\\ProgramData"),
        Path("C:\\Recovery"),
        Path("C:\\System Volume Information"),
        Path("C:\\$Recycle.Bin"),
    ]
elif _OS == "Linux" or _OS == "Darwin":
    _BLOCKED_PATHS = [
        Path("/etc"),
        Path("/sys"),
        Path("/proc"),
        Path("/dev"),
        Path("/boot"),
        Path("/root"),
        Path("/var/log"),
    ]

def _is_safe_path(target: Path) -> bool:
    """Check if path is in safe roots and not in blocked paths."""
    try:
        resolved = target.resolve()
        # Check against blocked paths
        for blocked in _BLOCKED_PATHS:
            if resolved == blocked.resolve() or resolved.is_relative_to(blocked.resolve()):
                return False
        # Check against safe roots
        return any(
            resolved == root.resolve() or resolved.is_relative_to(root.resolve())
            for root in _SAFE_ROOTS
        )
    except Exception:
        return False

def _get_desktop() -> Path:
    if _OS == "Linux":
        xdg = os.environ.get("XDG_DESKTOP_DIR", "")
        if xdg and Path(xdg).exists():
            return Path(xdg)
    return Path.home() / "Desktop"

def _get_downloads() -> Path:
    if _OS == "Linux":
        xdg = os.environ.get("XDG_DOWNLOAD_DIR", "")
        if xdg and Path(xdg).exists():
            return Path(xdg)
    return Path.home() / "Downloads"

def _get_documents() -> Path:
    if _OS == "Linux":
        xdg = os.environ.get("XDG_DOCUMENTS_DIR", "")
        if xdg and Path(xdg).exists():
            return Path(xdg)
    return Path.home() / "Documents"

def _get_pictures() -> Path:
    if _OS == "Linux":
        xdg = os.environ.get("XDG_PICTURES_DIR", "")
        if xdg and Path(xdg).exists():
            return Path(xdg)
    return Path.home() / "Pictures"

def _get_music() -> Path:
    if _OS == "Linux":
        xdg = os.environ.get("XDG_MUSIC_DIR", "")
        if xdg and Path(xdg).exists():
            return Path(xdg)
    return Path.home() / "Music"

def _get_videos() -> Path:
    if _OS == "Linux":
        xdg = os.environ.get("XDG_VIDEOS_DIR", "")
        if xdg and Path(xdg).exists():
            return Path(xdg)
    return Path.home() / "Videos"


def _resolve_path(raw: str) -> Path:
    shortcuts: dict[str, Path] = {
        "desktop":   _get_desktop(),
        "downloads": _get_downloads(),
        "documents": _get_documents(),
        "pictures":  _get_pictures(),
        "music":     _get_music(),
        "videos":    _get_videos(),
        "home":      Path.home(),
    }
    lower = raw.strip().lower()
    if lower in shortcuts:
        return shortcuts[lower]
    for name, path in shortcuts.items():
        if lower.startswith(name + "/") or lower.startswith(name + "\\"):
            return path / raw.strip()[len(name)+1:]
    # Check if it's an absolute path
    expanded = Path(raw).expanduser()
    if expanded.is_absolute():
        return expanded
    # Relative path — try to find on Desktop
    desktop = _get_desktop()
    candidate = desktop / raw.strip()
    if candidate.exists():
        return candidate
    # Fuzzy match: search Desktop for close folder name match
    best_match = _fuzzy_find(desktop, raw.strip())
    if best_match:
        return best_match
    # Folder index fallback — check all indexed folders
    idx = _get_folder_index()
    match = idx.quick_find(raw.strip(), threshold=0.7)
    if match:
        return Path(match)
    # Fallback: return Desktop-relative path anyway (caller handles not-found)
    return candidate


def _fuzzy_find(base_dir: Path, target: str) -> Path | None:
    """Search base_dir for a case-insensitive close match to target name."""
    target_lower = target.lower().replace(" ", "")
    best_score = 0
    best_path = None
    try:
        for item in base_dir.iterdir():
            name = item.name.lower().replace(" ", "")
            # Check if target is a substring of the item name or vice versa
            if target_lower in name or name in target_lower:
                score = len(set(target_lower) & set(name))
                if score > best_score:
                    best_score = score
                    best_path = item
            # Simple Levenshtein-like: count matching chars in order
            i = j = 0
            while i < len(target_lower) and j < len(name):
                if target_lower[i] == name[j]:
                    i += 1
                j += 1
            match_ratio = i / max(len(target_lower), 1)
            if match_ratio > 0.6 and match_ratio > best_score / max(len(target_lower), 1):
                best_score = int(match_ratio * len(target_lower))
                best_path = item
    except Exception:
        pass
    return best_path

def _format_size(b: int) -> str:
    for unit in ["B", "KB", "MB", "GB", "TB"]:
        if b < 1024:
            return f"{b:.1f} {unit}"
        b /= 1024
    return f"{b:.1f} TB"

def _safe_trash(target: Path) -> str:

    if _SEND2TRASH:
        send2trash.send2trash(str(target))
        return f"Moved to Trash: {target.name}"

    # Fallback: permanent delete when send2trash is not installed
    try:
        if target.is_dir():
            shutil.rmtree(str(target))
        else:
            target.unlink()
        return f"Permanently deleted: {target.name}"
    except Exception as e:
        return f"Could not delete: {e}"


def list_files(path: str = "desktop", show_hidden: bool = False) -> str:
    try:
        target = _resolve_path(path)
        if not _is_safe_path(target):
            return f"Access denied: {target}"
        if not target.exists():
            # Suggest similar folders on Desktop
            suggestions = _suggest_folders(path)
            if suggestions:
                return f"Path not found: '{path}'. Did you mean one of: {', '.join(suggestions)}?"
            return f"Path not found: {target}"
        if not target.is_dir():
            return f"Not a directory: {target.name}"

        items = []
        for item in sorted(target.iterdir()):
            if not show_hidden and item.name.startswith("."):
                continue
            if item.is_dir():
                items.append(f"📁 {item.name}/")
            else:
                size = _format_size(item.stat().st_size)
                items.append(f"📄 {item.name} ({size})")

        if not items:
            return f"Directory is empty: {target.name}/"

        return f"Contents of {target.name}/ ({len(items)} items):\n" + "\n".join(items)

    except PermissionError:
        return f"Permission denied: {path}"
    except Exception as e:
        return f"Error listing files: {e}"


def _suggest_folders(name: str) -> list[str]:
    """Find close-matching folder names across all indexed locations."""
    idx = _get_folder_index()
    results = idx.search(name, top_n=5)
    return [r[0] for r in results]


def create_file(path: str, name: str = "", content: str = "") -> str:
    try:
        base   = _resolve_path(path)
        target = (base / name) if name else base
        if not _is_safe_path(target):
            return f"Access denied: {target}"
        target.parent.mkdir(parents=True, exist_ok=True)

        suffix = target.suffix.lower()
        if suffix in (".xlsx", ".xls"):
            try:
                import openpyxl
                wb = openpyxl.Workbook()
                ws = wb.active
                ws.title = "Sheet1"
                if content:
                    lines = content.split('\n')
                    for r_idx, line in enumerate(lines, start=1):
                        cols = line.split(',')
                        for c_idx, val in enumerate(cols, start=1):
                            ws.cell(row=r_idx, column=c_idx, value=val.strip())
                wb.save(target)
                return f"Excel file created with valid structure: {target.name}"
            except ImportError:
                try:
                    import pandas as pd
                    df = pd.DataFrame()
                    df.to_excel(target, index=False)
                    return f"Excel file created with valid structure (pandas): {target.name}"
                except Exception as e:
                    pass
        elif suffix in (".docx", ".doc"):
            try:
                from docx import Document
                doc = Document()
                if content:
                    doc.add_paragraph(content)
                else:
                    doc.add_paragraph("")
                doc.save(target)
                return f"Word document created with valid structure: {target.name}"
            except ImportError:
                pass

        target.write_text(content, encoding="utf-8")
        _track_file(target)
        return f"File created: {target.name}"
    except Exception as e:
        return f"Could not create file: {e}"


def create_folder(path: str, name: str = "") -> str:
    try:
        base   = _resolve_path(path)
        target = (base / name) if name else base
        if not _is_safe_path(target):
            return f"Access denied: {target}"
        target.mkdir(parents=True, exist_ok=True)
        return f"Folder created: {target.name}"
    except Exception as e:
        return f"Could not create folder: {e}"


def delete_file(path: str, name: str = "") -> str:
    try:
        base   = _resolve_path(path)
        target = (base / name) if name else base
        if not _is_safe_path(target):
            return f"Access denied: {target}"
        if not target.exists():
            return f"Not found: {target.name}"

        # Güvenli dizin kontrolü — kritik kullanıcı klasörlerini koru
        protected = {
            _get_desktop(), _get_downloads(), _get_documents(),
            _get_pictures(), _get_music(), _get_videos(), Path.home()
        }
        if target.resolve() in {p.resolve() for p in protected}:
            return f"Protected directory, cannot delete: {target.name}"

        return _safe_trash(target)

    except PermissionError:
        return f"Permission denied: {path}"
    except Exception as e:
        return f"Could not delete: {e}"


def move_file(path: str, name: str = "", destination: str = "") -> str:
    try:
        base   = _resolve_path(path)
        src    = (base / name) if name else base
        dst    = _resolve_path(destination) if destination else None

        if not src.exists():
            return f"Source not found: {src.name}"
        if dst is None:
            return "No destination specified."
        if not _is_safe_path(src):
            return f"Access denied (source): {src}"
        if not _is_safe_path(dst):
            return f"Access denied (destination): {dst}"

        if dst.is_dir():
            dst = dst / src.name

        dst.parent.mkdir(parents=True, exist_ok=True)
        shutil.move(str(src), str(dst))
        return f"Moved: {src.name} -> {dst.parent.name}/"

    except Exception as e:
        return f"Could not move: {e}"


def copy_file(path: str, name: str = "", destination: str = "") -> str:
    try:
        base = _resolve_path(path)
        src  = (base / name) if name else base
        dst  = _resolve_path(destination) if destination else None

        if not src.exists():
            return f"Source not found: {src.name}"
        if dst is None:
            return "No destination specified."
        if not _is_safe_path(src):
            return f"Access denied (source): {src}"
        if not _is_safe_path(dst):
            return f"Access denied (destination): {dst}"

        if dst.is_dir():
            dst = dst / src.name

        dst.parent.mkdir(parents=True, exist_ok=True)

        if src.is_dir():
            shutil.copytree(str(src), str(dst))
        else:
            shutil.copy2(str(src), str(dst))

        return f"Copied: {src.name} -> {dst.parent.name}/"

    except Exception as e:
        return f"Could not copy: {e}"


def rename_file(path: str, name: str = "", new_name: str = "") -> str:
    try:
        base     = _resolve_path(path)
        target   = (base / name) if name else base
        if not _is_safe_path(target):
            return f"Access denied: {target}"
        if not target.exists():
            return f"Not found: {target.name}"
        if not new_name:
            return "No new name provided."

        new_path = target.parent / new_name
        if new_path.exists():
            return f"A file named '{new_name}' already exists here."

        target.rename(new_path)
        return f"Renamed: {target.name} -> {new_name}"

    except Exception as e:
        return f"Could not rename: {e}"


def read_file(path: str, name: str = "", max_chars: int = 4000) -> str:
    try:
        base   = _resolve_path(path)
        target = (base / name) if name else base
        if not _is_safe_path(target):
            return f"Access denied: {target}"
        if not target.exists():
            return f"File not found: {target.name}"
        if not target.is_file():
            return f"Not a file: {target.name}"

        content = target.read_text(encoding="utf-8", errors="ignore")
        if len(content) > max_chars:
            content = content[:max_chars] + f"\n\n[Truncated — {len(content)} total chars]"
        return content

    except Exception as e:
        return f"Could not read file: {e}"


def write_file(path: str, name: str = "", content: str = "",
               append: bool = False) -> str:
    try:
        base   = _resolve_path(path)
        target = (base / name) if name else base
        if not _is_safe_path(target):
            return f"Access denied: {target}"
        target.parent.mkdir(parents=True, exist_ok=True)
        mode = "a" if append else "w"
        with open(target, mode, encoding="utf-8") as f:
            f.write(content)
        _track_file(target)
        action = "Appended to" if append else "Written to"
        return f"{action}: {target.name}"
    except Exception as e:
        return f"Could not write file: {e}"


def find_files(name: str = "", extension: str = "",
               path: str = "home", max_results: int = 20) -> str:
    try:
        search_path = _resolve_path(path)
        if not _is_safe_path(search_path):
            return f"Access denied: {search_path}"
        if not search_path.exists():
            return f"Search path not found: {path}"

        results    = []
        dir_count  = 0
        max_dirs   = 500  # performans + güvenlik limiti

        for item in search_path.rglob("*"):
            if item.is_dir():
                dir_count += 1
                if dir_count > max_dirs:
                    break
                continue
            if not item.is_file():
                continue
            if extension and item.suffix.lower() != extension.lower():
                continue
            if name and name.lower() not in item.name.lower():
                continue
            size = _format_size(item.stat().st_size)
            results.append(f"📄 {item.name} ({size}) — {item.parent}")
            if len(results) >= max_results:
                break

        if not results:
            query = name or extension or "files"
            return f"No {query} found in {search_path.name}/"

        return f"Found {len(results)} file(s):\n" + "\n".join(results)

    except Exception as e:
        return f"Search error: {e}"


def get_largest_files(path: str = "downloads", count: int = 10) -> str:
    count = min(count, 50)  # maksimum 50
    try:
        search_path = _resolve_path(path)
        if not _is_safe_path(search_path):
            return f"Access denied: {search_path}"
        if not search_path.exists():
            return f"Path not found: {path}"

        files = []
        for item in search_path.rglob("*"):
            if item.is_file():
                try:
                    files.append((item.stat().st_size, item))
                except Exception:
                    continue

        files.sort(reverse=True)
        top = files[:count]

        if not top:
            return "No files found."

        lines = [f"Top {len(top)} largest files in {search_path.name}/:"]
        for size, f in top:
            lines.append(f"  {_format_size(size):>10}  {f.name}  ({f.parent})")

        return "\n".join(lines)

    except Exception as e:
        return f"Error: {e}"


def get_disk_usage(path: str = "home") -> str:
    try:
        target = _resolve_path(path)
        usage  = shutil.disk_usage(target)
        pct    = usage.used / usage.total * 100
        return (
            f"Disk usage ({target}):\n"
            f"  Total : {_format_size(usage.total)}\n"
            f"  Used  : {_format_size(usage.used)} ({pct:.1f}%)\n"
            f"  Free  : {_format_size(usage.free)}"
        )
    except Exception as e:
        return f"Could not get disk usage: {e}"


def organize_desktop() -> str:
    type_map = {
        "Images":    {".jpg", ".jpeg", ".png", ".gif", ".bmp", ".webp", ".svg", ".ico", ".heic"},
        "Documents": {".pdf", ".doc", ".docx", ".txt", ".xls", ".xlsx",
                      ".ppt", ".pptx", ".csv", ".odt", ".ods", ".odp"},
        "Videos":    {".mp4", ".avi", ".mkv", ".mov", ".wmv", ".flv", ".webm", ".m4v"},
        "Music":     {".mp3", ".wav", ".flac", ".aac", ".ogg", ".wma", ".m4a"},
        "Archives":  {".zip", ".rar", ".7z", ".tar", ".gz", ".bz2", ".xz"},
        "Code":      {".py", ".js", ".ts", ".html", ".css", ".json", ".xml",
                      ".cpp", ".java", ".cs", ".go", ".rs", ".sh"},
    }

    desktop = _get_desktop()
    moved, skipped = [], []

    try:
        for item in desktop.iterdir():
            # Klasörlere, gizli dosyalara ve organize klasörlerine dokunma
            if item.is_dir() or item.name.startswith("."):
                continue
            if item.name in {k for k in type_map}:
                continue

            ext        = item.suffix.lower()
            target_dir = desktop / "Others"
            for folder, exts in type_map.items():
                if ext in exts:
                    target_dir = desktop / folder
                    break

            target_dir.mkdir(exist_ok=True)
            new_path = target_dir / item.name

            if new_path.exists():
                skipped.append(item.name)
                continue

            shutil.move(str(item), str(new_path))
            moved.append(f"{item.name} -> {target_dir.name}/")

        result = f"Desktop organized: {len(moved)} files moved."
        if moved:
            preview = moved[:8]
            result += "\n" + "\n".join(preview)
            if len(moved) > 8:
                result += f"\n... and {len(moved) - 8} more."
        if skipped:
            result += f"\n{len(skipped)} file(s) skipped (name conflict)."
        return result

    except Exception as e:
        return f"Could not organize desktop: {e}"


def get_file_info(path: str, name: str = "") -> str:
    try:
        base   = _resolve_path(path)
        target = (base / name) if name else base
        if not _is_safe_path(target):
            return f"Access denied: {target}"
        if not target.exists():
            return f"Not found: {target.name}"

        stat = target.stat()
        info = {
            "Name":      target.name,
            "Type":      "Folder" if target.is_dir() else "File",
            "Size":      _format_size(stat.st_size),
            "Location":  str(target.parent),
            "Created":   datetime.fromtimestamp(stat.st_ctime).strftime("%Y-%m-%d %H:%M"),
            "Modified":  datetime.fromtimestamp(stat.st_mtime).strftime("%Y-%m-%d %H:%M"),
            "Extension": target.suffix or "—",
        }
        return "\n".join(f"  {k}: {v}" for k, v in info.items())

    except Exception as e:
        return f"Could not get file info: {e}"

# ---- New: RAG Semantic Search (Phase 5.1) ----

def _get_rag_collection():
    """Get or create ChromaDB collection for file content."""
    try:
        import posthog; posthog.capture = lambda *a, **kw: None
        import chromadb
        from chromadb.config import Settings as _CS
        db_path = Path(__file__).resolve().parent.parent / "memory" / "chroma_db"
        db_path.mkdir(parents=True, exist_ok=True)
        client = chromadb.PersistentClient(str(db_path), settings=_CS(anonymized_telemetry=False))
        return client.get_or_create_collection("file_index")
    except Exception:
        return None


def index_directory(path: str = "documents") -> str:
    """Index all text files in a directory into ChromaDB for semantic search."""
    try:
        target = _resolve_path(path)
        if not target.exists() or not target.is_dir():
            return f"Directory not found: {path}"

        collection = _get_rag_collection()
        if collection is None:
            return "ChromaDB not available. Install with: pip install chromadb"

        text_extensions = {
            ".py", ".js", ".ts", ".html", ".css", ".json", ".xml", ".md", ".txt",
            ".csv", ".yaml", ".yml", ".toml", ".ini", ".cfg", ".conf",
            ".java", ".cpp", ".c", ".h", ".hpp", ".rs", ".go", ".rb", ".php",
            ".sh", ".bat", ".ps1", ".sql", ".r", ".kt", ".swift",
        }

        indexed = 0
        skipped = 0
        for item in target.rglob("*"):
            if not item.is_file():
                continue
            if item.suffix.lower() not in text_extensions:
                skipped += 1
                continue
            try:
                content = item.read_text(encoding="utf-8", errors="replace")
                if len(content) < 20:
                    skipped += 1
                    continue
                rel_path = str(item.relative_to(target))
                collection.add(
                    documents=[content[:5000]],
                    metadatas=[{"source": rel_path, "path": str(item)}],
                    ids=[f"file_{abs(hash(str(item)))}"],
                )
                indexed += 1
            except Exception:
                skipped += 1

        return f"Indexed {indexed} files from {target.name}/ ({skipped} skipped). You can now use find_by_content to search."
    except Exception as e:
        return f"Indexing failed: {e}"


def find_by_content(query: str, path: str = "home") -> str:
    """Semantic search for files by content using ChromaDB RAG."""
    try:
        collection = _get_rag_collection()
        if collection is None:
            return "ChromaDB not available. Run 'index_directory' first or install chromadb."

        try:
            count = collection.count()
            if count == 0:
                return "No files indexed. Run 'index_directory' first to build the search index."
        except Exception:
            return "No files indexed. Run 'index_directory' first."

        results = collection.query(query_texts=[query], n_results=10)

        if not results or not results.get("documents") or not results["documents"][0]:
            return f"No semantic matches found for '{query}'. Try a different search term."

        output = f"🔍 Semantic search results for: '{query}'\n"
        for i, (doc, meta) in enumerate(zip(results["documents"][0], results["metadatas"][0])):
            source = meta.get("source", "unknown")
            score = results["distances"][0][i] if results.get("distances") else 0
            snippet = doc[:150].replace("\n", " ")
            output += f"\n[{i+1}] {source} (score: {score:.2f})\n    {snippet}..."

        return output
    except Exception as e:
        return f"Semantic search failed: {e}"


def search_code(query: str, path: str = "home") -> str:
    """Find code files by semantic meaning using ChromaDB."""
    try:
        target = _resolve_path(path)
        if not target.exists():
            return f"Path not found: {path}"

        collection = _get_rag_collection()
        if collection is not None:
            try:
                if collection.count() > 0:
                    results = collection.query(
                        query_texts=[f"code: {query}"],
                        n_results=10,
                        where={"source": {"$regex": "\\.(py|js|ts|java|cpp)$"}}
                    )
                    if results.get("documents") and results["documents"][0]:
                        output = f"🔍 Code search results for: '{query}'\n"
                        for doc, meta in zip(results["documents"][0], results["metadatas"][0]):
                            source = meta.get("source", "unknown")
                            snippet = doc[:200].replace("\n", " ")
                            output += f"\n📄 {source}\n  {snippet}..."
                        return output
            except Exception:
                pass

        # Fallback: grep-style search
        code_extensions = {".py", ".js", ".ts", ".java", ".cpp", ".c", ".h", ".go", ".rs", ".rb", ".php"}
        matches = []
        for item in target.rglob("*"):
            if item.is_file() and item.suffix.lower() in code_extensions:
                try:
                    content = item.read_text(encoding="utf-8", errors="replace")
                    if query.lower() in content.lower():
                        rel = str(item.relative_to(target))[:80]
                        matches.append(rel)
                except Exception:
                    pass
                if len(matches) >= 15:
                    break

        if matches:
            return f"Found {len(matches)} file(s) containing '{query}':\n" + "\n".join(matches)
        return f"No code files found containing '{query}'."
    except Exception as e:
        return f"Code search failed: {e}"


def _format_size(size: int) -> str:
    if _HUMANIZE:
        return humanize.naturalsize(size, binary=True)
    for unit in ("B", "KB", "MB", "GB", "TB"):
        if size < 1024:
            return f"{size:.1f} {unit}"
        size /= 1024
    return f"{size:.1f} PB"


def _detect_mime(path: Path) -> str:
    if _FILE_MAGIC:
        try:
            mime = magic.Magic(mime=True)
            return mime.from_file(str(path))
        except Exception:
            pass
    suffix_map = {
        ".txt": "text/plain", ".csv": "text/csv", ".json": "application/json",
        ".xml": "application/xml", ".html": "text/html", ".pdf": "application/pdf",
        ".png": "image/png", ".jpg": "image/jpeg", ".jpeg": "image/jpeg",
        ".gif": "image/gif", ".webp": "image/webp", ".mp3": "audio/mpeg",
        ".mp4": "video/mp4", ".zip": "application/zip", ".py": "text/x-python",
        ".js": "text/javascript", ".ts": "text/typescript", ".md": "text/markdown",
    }
    return suffix_map.get(path.suffix.lower(), "application/octet-stream")


def bulk_rename(path: Path, pattern: str, replacement: str, dry_run: bool = True) -> str:
    if not _WCMATCH:
        return "wcmatch not installed. Run: pip install wcmatch"
    results = []
    glob_re = re.escape(pattern).replace(r"\*", "(.*)")
    for item in path.rglob("*"):
        if item.is_file() and wc_glob.globmatch(item.name, pattern, flags=wc_glob.GLOBSTAR):
            m = re.match(glob_re, item.name)
            if not m:
                continue
            # Support {1} {2} or \1 \2 style replacements
            new_name = replacement
            for i, group in enumerate(m.groups(), 1):
                new_name = new_name.replace(f"{{{i}}}", group).replace(f"\\{i}", group)
            new_path = item.parent / new_name
            if new_path == item:
                continue
            if not dry_run:
                if new_path.exists():
                    new_path.unlink()
                item.rename(new_path)
            results.append(f"  {item.name} -> {new_name}")
    if not results:
        return f"No files matched pattern: {pattern}"
    summary = f"{'[DRY RUN] ' if dry_run else ''}Renamed {len(results)} files:\n"
    summary += "\n".join(results[:50])
    if len(results) > 50:
        summary += f"\n  ... and {len(results) - 50} more"
    return summary


def git_operation(path: Path, action: str, **kwargs) -> str:
    if not _GITPYTHON:
        return "GitPython not installed. Run: pip install gitpython"
    try:
        repo = git.Repo(path)
    except git.InvalidGitRepositoryError:
        return f"Not a git repository: {path}"

    try:
        if action == "status":
            return repo.git.status()
        elif action == "diff":
            return repo.git.diff()[:2000]
        elif action == "log":
            return repo.git.log(oneline=True, max_count=15)
        elif action == "commit":
            files = kwargs.get("files", ".")
            msg = kwargs.get("message", "Auto-commit from MARK XL")
            repo.index.add(files)
            repo.index.commit(msg)
            return f"Committed with message: {msg}"
        elif action == "add":
            repo.index.add(kwargs.get("files", "."))
            return f"Staged: {kwargs.get('files', '.')}"
        elif action == "pull":
            repo.remotes.origin.pull()
            return "Pulled latest from origin"
        elif action == "push":
            repo.remotes.origin.push()
            return "Pushed to origin"
        elif action == "branch":
            return "\n".join(f"  {'*' if b.name == repo.active_branch.name else ' '} {b.name}" for b in repo.branches)
        else:
            return f"Unknown git action: {action}. Use: status, diff, log, commit, add, pull, push, branch"
    except Exception as e:
        return f"Git {action} failed: {e}"


_watchdog_observers = {}


def file_watch_start(path: Path) -> str:
    if not _WATCHDOG:
        return "watchdog not installed. Run: pip install watchdog"

    class _AgentHandler(FileSystemEventHandler):
        def __init__(self, watch_path):
            self.watch_path = watch_path
        def on_modified(self, event):
            if not event.is_directory:
                print(f"[FileWatch] Modified: {event.src_path}")
        def on_created(self, event):
            if not event.is_directory:
                print(f"[FileWatch] Created: {event.src_path}")
        def on_deleted(self, event):
            if not event.is_directory:
                print(f"[FileWatch] Deleted: {event.src_path}")

    observer = Observer()
    handler = _AgentHandler(path)
    observer.schedule(handler, str(path), recursive=True)
    observer.start()
    _watchdog_observers[str(path)] = observer
    return f"Watching: {path}"


def file_watch_stop(path: Path) -> str:
    key = str(path)
    if key in _watchdog_observers:
        _watchdog_observers[key].stop()
        _watchdog_observers[key].join()
        del _watchdog_observers[key]
        return f"Stopped watching: {path}"
    return f"No watcher found for: {path}"


def file_controller(
    parameters: dict = None,
    response=None,
    player=None,
    session_memory=None,
) -> str:
    params = parameters or {}
    action = params.get("action", "").lower().strip()
    path   = params.get("path", "desktop")
    name   = params.get("name", "")

    if player:
        player.write_log(f"[file] {action} {name or path}")

    try:
        from memory.memory_manager import update_memory
        from core.agent_tracker import log_activity

        result = ""
        if action == "list":
            result = list_files(path)

        elif action == "create_file":
            result = create_file(path, name=name, content=params.get("content", ""))
            log_activity("file_controller", f"create_file: {name}", "success" if "Created" in result else "fail")
            try:
                resolved = _resolve_path(path) / name if name else _resolve_path(path)
                update_memory({"projects": {f"file_{int(time.time())}": {"value": str(resolved)}}})
            except Exception:
                pass

        elif action == "create_folder":
            result = create_folder(path, name=name)
            log_activity("file_controller", f"create_folder: {name or path}", "success" if "Created" in result else "fail")
            try:
                resolved = _resolve_path(path) / name if name else _resolve_path(path)
                update_memory({"projects": {f"folder_{int(time.time())}": {"value": str(resolved)}}})
            except Exception:
                pass

        elif action == "delete":
            result = delete_file(path, name=name)

        elif action == "move":
            result = move_file(path, name=name, destination=params.get("destination", ""))

        elif action == "copy":
            result = copy_file(path, name=name, destination=params.get("destination", ""))

        elif action == "rename":
            result = rename_file(path, name=name, new_name=params.get("new_name", ""))

        elif action == "read":
            result = read_file(path, name=name)

        elif action == "write":
            result = write_file(
                path, name=name,
                content=params.get("content", ""),
                append=params.get("append", False)
            )

        elif action == "find":
            cached = _find_recent(name or params.get("name", ""))
            if cached:
                result = f"Found: {cached}\n"
            result = (result or "") + find_files(
                name=name or params.get("name", ""),
                extension=params.get("extension", ""),
                path=path,
                max_results=min(int(params.get("max_results", 20)), 50),
            )

        elif action == "find_by_content":
            result = find_by_content(query=params.get("query", name), path=path)

        elif action == "search_code":
            result = search_code(query=params.get("query", name), path=path)

        elif action == "find_folder":
            query = name or params.get("query", "")
            if not query:
                result = "Provide a folder name to search."
            else:
                idx = _get_folder_index()
                matches = idx.search(query, top_n=5)
                if not matches:
                    result = f"No folders found matching '{query}'."
                elif len(matches) == 1 and matches[0][2] >= 0.85:
                    result = f"Opened: {matches[0][1]}"
                    try:
                        os.startfile(matches[0][1])
                    except Exception:
                        pass
                else:
                    lines = [f"Found {len(matches)} folders matching '{query}':"]
                    for i, (name, path, score) in enumerate(matches, 1):
                        pct = int(score * 100)
                        lines.append(f"  {i}. {name}  ({path})  [{pct}%]")
                    result = "\n".join(lines)

        elif action == "index_directory":
            result = index_directory(path=path)

        elif action == "largest":
            result = get_largest_files(
                path=path,
                count=int(params.get("count", 10)),
            )

        elif action == "disk_usage":
            result = get_disk_usage(path)

        elif action == "organize_desktop":
            result = organize_desktop()

        elif action == "info":
            result = get_file_info(path, name=name)

        elif action in ("run", "execute", "open_file"):
            import subprocess, sys
            target_name = name or params.get("name", "")
            target_path = params.get("path", "")
            try:
                resolved = _resolve_path(target_path) / target_name if target_name else _resolve_path(target_path)
            except Exception:
                resolved = None
            if not resolved or not resolved.exists():
                cached = _find_recent(target_name or target_path)
                if cached:
                    resolved = Path(cached)
            if resolved and resolved.exists():
                if resolved.suffix == ".py":
                    proc = subprocess.run([sys.executable, str(resolved)], capture_output=True, text=True, timeout=30)
                    out = proc.stdout.strip() or proc.stderr.strip() or "No output"
                    return f"Ran {resolved.name}. Output:\n{out[:500]}"
                else:
                    os.startfile(str(resolved))
                    return f"Opened: {resolved.name}"
            else:
                return f"File not found: {target_name or target_path}"

        elif action == "bulk_rename":
            pattern = params.get("pattern", "")
            replacement = params.get("replacement", "")
            dry_run = params.get("dry_run", True)
            if not pattern or replacement is None:
                return "bulk_rename requires 'pattern' and 'replacement' parameters"
            result = bulk_rename(path, pattern, replacement, dry_run=dry_run)

        elif action == "mime":
            resolved = _resolve_path(name or path)
            if resolved and resolved.exists():
                detected = _detect_mime(resolved)
                result = f"{resolved.name}: {detected}"
            else:
                result = "File not found"

        elif action == "git":
            git_action = params.get("git_action", "status")
            git_msg = params.get("message", "Auto-commit from MARK XL")
            git_files = params.get("files", ".")
            result = git_operation(path, git_action, message=git_msg, files=git_files)

        elif action == "watch_start":
            result = file_watch_start(path)

        elif action == "watch_stop":
            result = file_watch_stop(path)

        else:
            return f"Unknown action: '{action}'"

        log_activity("file_controller", action, "success", detail=name or path)
        return result

    except Exception as e:
        return f"File controller error ({action}): {e}"
